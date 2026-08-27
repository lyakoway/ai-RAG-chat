"""Генерирует демо-документы в backend/samples для проверки RAG — RU и EN.

Запуск:
    cd backend
    .venv/bin/python scripts/make_samples.py

Создаёт 6 файлов: Политика_удалённой_работы.docx, Тарифы_и_скидки.xlsx,
Руководство_пользователя.pdf и их английские версии (Remote_Work_Policy.docx,
Pricing_and_Discounts.xlsx, User_Guide.pdf).

PDF создаётся, если установлен reportlab (`pip install reportlab`);
иначе шаг PDF пропускается — docx/xlsx создаются всегда.
"""
from __future__ import annotations

from pathlib import Path

SAMPLES = Path(__file__).resolve().parent.parent / "samples"

# Содержимое обоих языков: (заголовок, [(подзаголовок, абзац), ...]).
POLICY = {
    "ru": (
        "Политика удалённой работы",
        [
            ("1. Общие положения",
             "Сотрудники могут работать удалённо до 3 дней в неделю по согласованию "
             "с непосредственным руководителем. Полностью удалённый формат согласуется "
             "с директором департамента."),
            ("2. Рабочее время",
             "Рабочий день длится 8 часов. Обязательные часы присутствия онлайн — "
             "с 11:00 до 16:00. В это время сотрудник должен быть доступен в корпоративном мессенджере."),
            ("3. Оборудование",
             "Компания предоставляет ноутбук и компенсирует до 3000 руб./мес. на интернет. "
             "Заявка на компенсацию подаётся через портал до 5 числа месяца."),
            ("4. Отпуск",
             "Ежегодный оплачиваемый отпуск — 28 календарных дней. Заявление подаётся "
             "не позднее чем за 14 дней до начала отпуска."),
        ],
    ),
    "en": (
        "Remote Work Policy",
        [
            ("1. General Provisions",
             "Employees may work remotely up to 3 days per week upon agreement with "
             "their direct manager. A fully remote arrangement is approved by the "
             "department director."),
            ("2. Working Hours",
             "The workday lasts 8 hours. Mandatory online availability hours are "
             "11:00–16:00. During this time the employee must be reachable in the "
             "corporate messenger."),
            ("3. Equipment",
             "The company provides a laptop and reimburses up to 3,000 RUB/month "
             "for internet. Reimbursement requests are submitted via the portal "
             "no later than the 5th day of the month."),
            ("4. Vacation",
             "Annual paid leave is 28 calendar days. The request must be submitted "
             "at least 14 days before the vacation starts."),
        ],
    ),
}

PRICING = {
    "ru": {
        "Тарифы": {
            "head": ["Тариф", "Цена, руб/мес", "Пользователей", "Поддержка"],
            "rows": [
                ["Старт", 990, 5, "Email"],
                ["Бизнес", 2990, 25, "Email + чат"],
                ["Корпоративный", 9990, "без лимита", "Персональный менеджер"],
            ],
        },
        "Скидки": {
            "head": ["Условие", "Скидка"],
            "rows": [
                ["Оплата за год", "20%"],
                ["Некоммерческие организации", "30%"],
            ],
        },
    },
    "en": {
        "Pricing": {
            "head": ["Plan", "Price, RUB/mo", "Users", "Support"],
            "rows": [
                ["Start", 990, 5, "Email"],
                ["Business", 2990, 25, "Email + chat"],
                ["Enterprise", 9990, "unlimited", "Dedicated manager"],
            ],
        },
        "Discounts": {
            "head": ["Condition", "Discount"],
            "rows": [
                ["Annual payment", "20%"],
                ["Non-profit organizations", "30%"],
            ],
        },
    },
}

GUIDE = {
    "ru": {
        "filename": "Руководство_пользователя.pdf",
        "pages": [
            ("Руководство пользователя — Стр. 1", [
                "Раздел 1. Начало работы.",
                "Зарегистрируйтесь на портале и подтвердите email.",
                "Первый вход выполняется по ссылке из письма.",
                "Пароль должен содержать не менее 8 символов.",
            ]),
            ("Руководство пользователя — Стр. 2", [
                "Раздел 2. Импорт данных.",
                "Поддерживаются форматы CSV и XLSX до 50 МБ.",
                "Импорт запускается на вкладке 'Данные' -> 'Загрузить'.",
                "Ошибки импорта отображаются в журнале операций.",
            ]),
            ("Руководство пользователя — Стр. 3", [
                "Раздел 3. Безопасность.",
                "Включите двухфакторную аутентификацию в настройках.",
                "Сессия завершается автоматически через 30 минут простоя.",
                "Экспорт данных доступен только администраторам.",
            ]),
        ],
    },
    "en": {
        "filename": "User_Guide.pdf",
        "pages": [
            ("User Guide — Page 1", [
                "Section 1. Getting started.",
                "Register on the portal and confirm your email.",
                "The first login is done via the link from the email.",
                "The password must be at least 8 characters.",
            ]),
            ("User Guide — Page 2", [
                "Section 2. Data import.",
                "CSV and XLSX formats up to 50 MB are supported.",
                "Import is launched on the 'Data' -> 'Upload' tab.",
                "Import errors are shown in the operations log.",
            ]),
            ("User Guide — Page 3", [
                "Section 3. Security.",
                "Enable two-factor authentication in the settings.",
                "The session ends automatically after 30 minutes of inactivity.",
                "Data export is available to administrators only.",
            ]),
        ],
    },
}

DOCX_NAMES = {"ru": "Политика_удалённой_работы.docx", "en": "Remote_Work_Policy.docx"}
XLSX_NAMES = {"ru": "Тарифы_и_скидки.xlsx", "en": "Pricing_and_Discounts.xlsx"}


def make_docx(lang: str) -> None:
    from docx import Document

    title, sections = POLICY[lang]
    doc = Document()
    doc.add_heading(title, 0)
    for heading, paragraph in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(paragraph)
    doc.save(SAMPLES / DOCX_NAMES[lang])


def make_xlsx(lang: str) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    first = True
    for sheet_name, sheet in PRICING[lang].items():
        ws = wb.active if first else wb.create_sheet(sheet_name)
        if first:
            ws.title = sheet_name
            first = False
        ws.append(sheet["head"])
        for row in sheet["rows"]:
            ws.append(row)
    wb.save(SAMPLES / XLSX_NAMES[lang])


def make_pdf(lang: str) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return False

    guide = GUIDE[lang]
    # Кириллический шрифт нужен только для RU; для EN хватает Helvetica.
    font = "Helvetica"
    if lang == "ru":
        for path in ["/System/Library/Fonts/Supplemental/Arial.ttf",
                     "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]:
            if Path(path).exists():
                pdfmetrics.registerFont(TTFont("Body", path))
                font = "Body"
                break

    c = canvas.Canvas(str(SAMPLES / guide["filename"]), pagesize=A4)
    w, h = A4

    def page(title: str, lines: list[str]) -> None:
        c.setFont(font, 18)
        c.drawString(2 * cm, h - 3 * cm, title)
        c.setFont(font, 11)
        y = h - 4.5 * cm
        for line in lines:
            c.drawString(2 * cm, y, line)
            y -= 0.8 * cm
        c.showPage()

    for title, lines in guide["pages"]:
        page(title, lines)
    c.save()
    return True


def main() -> None:
    SAMPLES.mkdir(exist_ok=True)
    made: list[str] = []
    for lang in ("ru", "en"):
        make_docx(lang)
        make_xlsx(lang)
        made += [DOCX_NAMES[lang], XLSX_NAMES[lang]]
        if make_pdf(lang):
            made.append(GUIDE[lang]["filename"])
    print("Создано в", SAMPLES)
    for m in made:
        print("  •", m)
    if not Path(SAMPLES / GUIDE["ru"]["filename"]).exists():
        print("  (PDF пропущен: установите reportlab — pip install reportlab)")


if __name__ == "__main__":
    main()
