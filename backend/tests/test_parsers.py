"""Тесты парсеров PDF / DOCX / XLSX на лету создаваемых файлах."""
import pytest

from app.parsers.base import ParseError, page_count, parse_document

DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _make_docx(tmp_path):
    from docx import Document

    p = tmp_path / "policy.docx"
    doc = Document()
    doc.add_heading("Политика", 0)
    doc.add_paragraph("Текст политики удалённой работы.")
    doc.save(p)
    return p


def _make_xlsx(tmp_path):
    from openpyxl import Workbook

    p = tmp_path / "pricing.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Тарифы"
    ws.append(["Тариф", "Цена"])
    ws.append(["Старт", 990])
    wb.save(p)
    return p


def _make_pdf(tmp_path):
    reportlab = pytest.importorskip("reportlab")
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    p = tmp_path / "guide.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, "Section 1. Getting started")
    c.showPage()
    c.drawString(72, 720, "Section 2. Security")
    c.save()
    return p


def test_parse_docx(tmp_path):
    segs = parse_document(_make_docx(tmp_path), DOCX_TYPE)
    assert segs and "Текст политики" in segs[0].text
    assert page_count(segs) >= 1


def test_parse_xlsx_with_sheet_labels(tmp_path):
    segs = parse_document(_make_xlsx(tmp_path), XLSX_TYPE)
    assert any("Старт" in s.text for s in segs)
    assert any(s.label == "Тарифы" for s in segs)


def test_parse_pdf_pages(tmp_path):
    segs = parse_document(_make_pdf(tmp_path), "application/pdf")
    assert len(segs) == 2
    assert segs[0].page == 1 and segs[1].page == 2
    assert page_count(segs) == 2


def test_unsupported_type_raises(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("hi", encoding="utf-8")
    with pytest.raises(ParseError):
        parse_document(p, "text/plain")
